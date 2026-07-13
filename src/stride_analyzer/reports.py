"""Geração de relatórios JSON, PDF e DOCX."""

from __future__ import annotations

import html
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

FONT_REGULAR = "Helvetica"
FONT_BOLD = "Helvetica-Bold"
FONT_MONO = "Courier"


def _escape_pdf_text(text: str | None) -> str:
    return html.escape(text or "").replace("\n", "<br/>")


def _pdf_styles() -> dict[str, ParagraphStyle]:
    return {
        "title": ParagraphStyle(
            "Title",
            fontName=FONT_BOLD,
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            fontName=FONT_REGULAR,
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
            textColor=colors.grey,
            spaceAfter=24,
        ),
        "heading": ParagraphStyle(
            "Heading",
            fontName=FONT_BOLD,
            fontSize=13,
            leading=16,
            spaceBefore=12,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "Body",
            fontName=FONT_REGULAR,
            fontSize=10,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        ),
        "mono": ParagraphStyle(
            "Mono",
            fontName=FONT_MONO,
            fontSize=9,
            leading=12,
            spaceAfter=6,
        ),
    }


def gerar_pdf_relatorio(resultados: list[dict], pdf_path: Path) -> Path:
    """Gera PDF consolidado com análises STRIDE."""
    styles = _pdf_styles()
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Relatório STRIDE - Diagramas de Arquitetura",
        author="TechChallenger",
    )

    story = []
    agora = datetime.now().strftime("%d/%m/%Y %H:%M")
    story.append(Paragraph("Relatório de Análise STRIDE", styles["title"]))
    story.append(
        Paragraph(
            f"Diagramas de arquitetura validados | Gerado em {agora}",
            styles["subtitle"],
        )
    )
    story.append(Spacer(1, 0.5 * cm))

    for idx, item in enumerate(resultados, start=1):
        nome = item["arquivo"]
        analise = item["analise"]
        img_path = Path(item["caminho"])

        story.append(Paragraph(f"Diagrama {idx}: {_escape_pdf_text(nome)}", styles["heading"]))

        if img_path.exists():
            try:
                img = RLImage(str(img_path), width=14 * cm, height=8 * cm, kind="proportional")
                story.append(img)
                story.append(Spacer(1, 0.3 * cm))
            except Exception:
                story.append(
                    Paragraph(
                        f"[Miniatura indisponível: {_escape_pdf_text(str(img_path))}]",
                        styles["mono"],
                    )
                )

        for paragrafo in (analise or "").split("\n\n"):
            paragrafo = paragrafo.strip()
            if paragrafo:
                story.append(Paragraph(_escape_pdf_text(paragrafo), styles["body"]))

        if idx < len(resultados):
            story.append(PageBreak())

    doc.build(story)
    return pdf_path


def _configurar_fonte_docx(
    documento: Document, nome_fonte: str = "Arial", tamanho: int = 11
) -> None:
    estilo = documento.styles["Normal"]
    fonte = estilo.font
    fonte.name = nome_fonte
    fonte.size = Pt(tamanho)
    estilo.element.rPr.rFonts.set(qn("w:eastAsia"), nome_fonte)


def gerar_docx_relatorio(resultados: list[dict], docx_path: Path) -> Path:
    """Gera relatório DOCX consolidado."""
    doc = Document()
    _configurar_fonte_docx(doc)

    titulo = doc.add_heading("Relatório de Análise STRIDE", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    agora = datetime.now().strftime("%d/%m/%Y %H:%M")
    subtitulo = doc.add_paragraph(
        f"Diagramas de arquitetura validados | Gerado em {agora}"
    )
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, item in enumerate(resultados, start=1):
        nome = item["arquivo"]
        analise = item["analise"]
        img_path = Path(item["caminho"])

        doc.add_heading(f"Diagrama {idx}: {nome}", level=1)

        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(5.5))
            except Exception:
                doc.add_paragraph(f"[Imagem indisponível: {img_path}]")

        for paragrafo in (analise or "").split("\n\n"):
            paragrafo = paragrafo.strip()
            if paragrafo:
                p = doc.add_paragraph(paragrafo)
                p.paragraph_format.space_after = Pt(6)

        if idx < len(resultados):
            doc.add_page_break()

    doc.save(str(docx_path))
    return docx_path


def salvar_relatorios(
    resultados: list[dict],
    relatorio_dir: Path,
    timestamp: str | None = None,
) -> dict[str, Path]:
    """Persiste JSON, PDF e DOCX na pasta relatorio/."""
    relatorio_dir.mkdir(parents=True, exist_ok=True)
    if timestamp is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    json_path = relatorio_dir / f"analise_stride_{timestamp}.json"
    pdf_path = relatorio_dir / f"relatorio_stride_{timestamp}.pdf"
    docx_path = relatorio_dir / f"relatorio_stride_{timestamp}.docx"

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(resultados, f, ensure_ascii=False, indent=2)

    gerar_pdf_relatorio(resultados, pdf_path)
    gerar_docx_relatorio(resultados, docx_path)

    return {"json": json_path, "pdf": pdf_path, "docx": docx_path}
